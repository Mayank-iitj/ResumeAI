"""
DOCX Parser Module
Extracts text and tables from DOCX resumes using python-docx
"""
from docx import Document
from pathlib import Path
from typing import Dict, List
import logging

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parse DOCX files and extract text content"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
    
    def parse(self, file_path: str) -> Dict[str, any]:
        """
        Parse DOCX file and extract text, tables, and metadata
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing extracted text, tables, and metadata
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            result = {
                'text': '',
                'tables': [],
                'paragraphs': 0,
                'metadata': {},
                'file_path': str(file_path)
            }
            
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            result['text'] = '\n'.join(paragraphs)
            result['paragraphs'] = len(paragraphs)
            
            # Extract tables
            tables_data = []
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append({
                    'table_num': table_num,
                    'data': table_data
                })
            
            result['tables'] = tables_data
            
            # Extract metadata
            core_properties = doc.core_properties
            result['metadata'] = {
                'author': core_properties.author,
                'created': str(core_properties.created) if core_properties.created else None,
                'modified': str(core_properties.modified) if core_properties.modified else None,
                'title': core_properties.title,
                'subject': core_properties.subject
            }
            
            logger.info(f"Successfully parsed DOCX: {file_path.name} ({result['paragraphs']} paragraphs)")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    def extract_text_only(self, file_path: str) -> str:
        """
        Quick extraction of text only (no tables)
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        """
        result = self.parse(file_path)
        return result['text']
    
    def is_valid_docx(self, file_path: str) -> bool:
        """
        Check if file is a valid DOCX
        
        Args:
            file_path: Path to check
            
        Returns:
            True if valid DOCX, False otherwise
        """
        try:
            doc = Document(file_path)
            return True
        except:
            return False
